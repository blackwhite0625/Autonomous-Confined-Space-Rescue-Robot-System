"""
產生專題內容（一、摘要 ～ 七、指導內容）的 .docx 檔案。
僅 3 張圖：(一)系統架構圖  (二)搜救流程圖  (三)車子實體照片
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── 全域樣式 ──
style = doc.styles['Normal']
font = style.font
font.name = '新細明體'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

def heading(text, bold=True, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = '新細明體'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = '新細明體'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = '新細明體'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')

# ================================================================
# 一、摘要
# ================================================================
heading('一、摘要')
body(
    '本研究主要開發一套「具生命跡象偵測之搜救機器人系統」。'
    '以樹莓派5B（Raspberry Pi 5B, 8GB）為核心控制平台，搭配Hailo-8 AI加速器（PCIe AI Kit）進行即時人體偵測與姿態辨識，'
    '結合超聲波避障、音訊分析（語音活動偵測、呼救聲辨識、敲擊聲偵測）、'
    '遠端光體積描記術（rPPG）心率估算等多模態感測技術，'
    '實現災後環境中自主巡邏、受困者偵測、生命跡象評估及遠端通報等功能。'
    '系統透過Flask網頁介面提供即時影像監控與手動遙控能力，'
    '並整合Telegram Bot推播實現緊急事件通報。',
    indent=True
)

# ================================================================
# 二、研究動機與研究問題
# ================================================================
heading('二、研究動機與研究問題')
body(
    '近年來全球天然災害與意外事故頻繁發生。根據內政部消防署統計資料指出，'
    '臺灣每年因地震、颱風、土石流等天然災害造成之建築物倒塌與人員受困事件時有所聞。'
    '在災後搜救行動中，黃金72小時為救援關鍵時間窗口，'
    '然而倒塌建物內部結構複雜、空間狹小、粉塵瀰漫，'
    '搜救人員進入現場不僅耗時且具有高度危險性。',
    indent=True
)
body(
    '因此，本研究提出以自主搜救機器人作為搜救人員的前導偵察工具，'
    '透過搭載AI人體偵測、姿態辨識、音訊分析及生命跡象感測等多模態技術，'
    '使機器人能夠在狹小且危險的空間中自主巡邏，'
    '即時偵測受困者並評估其生命跡象狀態，'
    '將偵測結果透過無線網路即時回傳至後方指揮中心，'
    '大幅縮短搜索時間並降低搜救人員的風險。',
    indent=True
)
body(
    '本研究欲解決之問題包含：'
    '（1）如何在資源受限的嵌入式平台上實現即時AI人體偵測與姿態辨識；'
    '（2）如何整合視覺、音訊、距離等多模態感測器進行受困者風險評估；'
    '（3）如何設計可靠的自主巡邏策略使機器人有效探索未知環境；'
    '（4）如何透過非接觸式方法遠端偵測受困者的生命跡象。',
    indent=True
)

# ================================================================
# 三、文獻回顧與探討
# ================================================================
heading('三、文獻回顧與探討')

heading('3-1 YOLOv8人體偵測與姿態辨識（參考文獻[1][2]）', size=12)  # YOLOv8 文件+GitHub
body(
    'YOLOv8為Ultralytics於2023年發布之最新物件偵測模型，'
    '採用anchor-free設計與C2f模組，在維持高精度的同時顯著提升推論速度。'
    '本系統採用YOLOv8n（Nano）版本進行人體偵測，以及YOLOv8s-pose版本進行17個關鍵點之人體姿態估計。'
    '透過骨架關鍵點的幾何分析（如肩-臀連線角度、骨架寬高比、邊界框長寬比），'
    '可判斷人員是否處於倒地、蜷縮、疑似不適等異常姿態，'
    '作為受困者偵測的重要視覺線索。',
    indent=True
)

heading('3-2 Hailo-8 AI加速器（參考文獻[3][4]）', size=12)  # Hailo + RPi5
body(
    'Hailo-8為以色列Hailo公司開發之邊緣AI處理器，'
    '提供高達26 TOPS（Tera Operations Per Second）的運算能力，'
    '功耗僅2.5W。透過Raspberry Pi 5的PCIe介面連接，'
    '可將YOLOv8模型編譯為HEF格式於NPU上執行，'
    '實現每秒15-20幀的即時推論，'
    '相較於純CPU推論（ONNX Runtime）速度提升約3-5倍。',
    indent=True
)

heading('3-3 遠端光體積描記術rPPG（參考文獻[5]）', size=12)  # rPPG 技術介紹
body(
    '遠端光體積描記術（Remote Photoplethysmography, rPPG）是一種非接觸式生命跡象偵測技術，'
    '透過普通RGB攝影機捕捉人臉皮膚表面因心臟搏動而產生的微小顏色變化，'
    '經由訊號處理萃取出心率資訊。'
    '本系統採用綠色通道（Green Channel）平均值法，'
    '搭配帶通濾波器（0.7-3.5Hz，對應42-210 BPM）與FFT頻譜分析，'
    '從臉部前額區域的ROI中估計受困者心率。'
    '此方法無需任何穿戴式裝置即可遠端評估受困者是否具有生命跡象，'
    '對於災後搜救場景具有重要應用價值。',
    indent=True
)

heading('3-4 麥克納姆輪全向移動平台（參考文獻[6]）', size=12)  # Mecanum 教學
body(
    '麥克納姆輪（Mecanum Wheel）為一種特殊設計的全向輪，'
    '輪緣上安裝有45度角的被動滾子，'
    '透過四輪獨立控制可實現前進、後退、橫移、斜移及原地旋轉等全向運動。'
    '其運動學公式為：'
    'FL = y + x + r、FR = y - x - r、RL = y - x + r、RR = y + x - r，'
    '其中x為橫移分量、y為前後分量、r為旋轉分量。'
    '本系統搭載4組直流馬達搭配2塊L298N驅動板，'
    '可在狹小空間中靈活移動，特別適合災後建物內部的複雜地形。'
    '如圖(三)所示為本系統機器人實體照片。',
    indent=True
)
caption('圖(三) 搜救機器人實體照片')
body('（請於此處插入車子實體照片）', indent=True)

heading('3-5 多模態融合與決策（參考文獻[7][8][9]）', size=12)  # Flask + gpiozero + 消防署
body(
    '本系統採用加權融合方法，整合視覺偵測、姿態分析、音訊偵測、'
    '互動回應、距離測量及生命跡象等六種感測模態，'
    '計算受困者風險分數（VictimScore）。'
    '融合公式為：VictimScore = 0.40×Person + 0.18×Pose + 0.18×Audio + 0.09×Motion + 0.05×Distance + 0.10×VitalSigns。'
    '根據分數高低劃分為LOW（<0.30）、SUSPECT（0.30-0.60）、HIGH（≥0.60）三個風險等級，'
    '驅動7階段任務狀態機進行對應的搜救行動。',
    indent=True
)

# ================================================================
# 四、研究方法與步驟
# ================================================================
heading('四、研究方法與步驟')

body(
    '研究方法系統架構如圖(一)所示，本系統以Raspberry Pi 5B為核心，'
    '透過USB介面連接攝影機、麥克風與喇叭，'
    '以GPIO控制馬達驅動板（L298N×2）、舵機雲台（SG90×2）與超聲波感測器（HC-SR04），'
    '並透過PCIe介面連接Hailo-8 NPU進行AI加速推論。'
    '系統透過Flask網頁伺服器提供即時MJPEG影像串流與RESTful API控制介面，'
    '使用者可透過瀏覽器進行遠端監控與操作。',
    indent=True
)
caption('圖(一) 系統架構圖')
body('（請於此處插入系統架構圖截圖）', indent=True)

heading('系統軟體架構', size=12, bold=True)
body(
    '軟體架構採用多執行緒設計，共有4條背景執行緒同時運行：'
    '（1）AI偵測迴圈：負責即時影像推論、人體偵測、姿態分析、rPPG心率估算；'
    '（2）音訊偵測迴圈：負責麥克風監聽、語音活動偵測、呼救關鍵字辨識；'
    '（3）超聲波測距迴圈：負責距離量測與安全煞車；'
    '（4）搜索巡邏迴圈：負責自主巡邏策略執行與避障控制。',
    indent=True
)

heading('7階段任務狀態機', size=12, bold=True)
body(
    '本系統設計7階段任務狀態機控制搜救流程，如圖(二)所示：'
    '（1）STANDBY待命：系統初始化完成，等待啟動；'
    '（2）SEARCH搜索：自主巡邏探索環境，舵機攝影機左右掃描偵測人員；'
    '（3）ANOMALY異常：VictimScore≥0.30，疑似發現受困者，持續觀察確認；'
    '（4）LOCK_ON鎖定：VictimScore≥0.60，確認高風險目標，機器人漸進靠近；'
    '（5）INQUIRY問詢：距離受困者≤30cm，啟動語音互動詢問傷勢狀況；'
    '（6）CONFIRM確認：彙整所有感測資料，最終確認是否需要救援；'
    '（7）REPORT回報：發送Telegram通知、播放警報音、記錄事件截圖與座標資訊。',
    indent=True
)
caption('圖(二) 7階段任務狀態機搜救流程圖')
body('（請於此處插入搜救流程圖截圖）', indent=True)

heading('自主巡邏策略', size=12, bold=True)
body(
    '本系統提供三種搜索模式：'
    '（1）模式D靜止掃描：機器人停留原地，舵機雲台左右掃描（±50°）偵測周圍人員；'
    '（2）模式E智慧巡邏：採用走停掃描策略，前進3秒後停車進行攝影機掃描5秒，遇障礙時後退轉向；'
    '（3）模式F掃描巡邏：利用麥克納姆輪原地旋轉特性，結合超聲波感測器掃描多個方向（±60°、±30°、0°共5個方向）'
    '的障礙距離，選擇最開闊且偏向未探索區域的方向前進，有效避免重複搜索與牆角卡死問題。'
    '系統同時搭配熱區記憶地圖（Heat Map），以航位推算記錄機器人行進軌跡與已探索區域。',
    indent=True
)

heading('人機互動模組（HRI）', size=12, bold=True)
body(
    '當機器人靠近疑似受困者時，啟動主動語音互動模組。'
    '系統透過Google Text-to-Speech發出中文問詢語音（如「有人被困嗎？請回應或發出聲音」），'
    '再透過USB麥克風收音並以Google Speech Recognition進行語音辨識，'
    '分析受困者回應中是否包含求救關鍵字（如「救命」、「幫忙」、「help」等）。'
    '若偵測到明確求救語音，系統立即進入確認與通報階段。',
    indent=True
)

heading('學習步驟', size=12, bold=True)
body(
    '(1) 學習樹莓派Linux系統操作與GPIO控制：'
    '本系統以Raspberry Pi 5B作為開發核心，首先學習Linux（Debian Bookworm）系統操作、'
    'Python程式撰寫、GPIO腳位控制，以及gpiozero函式庫的PWM輸出與感測器讀取。',
    indent=True
)
body(
    '(2) 學習YOLOv8模型部署與Hailo NPU加速：'
    '學習ultralytics框架的模型訓練與推論，'
    '以及Hailo SDK的HEF模型編譯與HailoRT推論引擎的使用，'
    '實現ONNX與Hailo雙後端的無縫切換。',
    indent=True
)
body(
    '(3) 學習Flask網頁伺服器與前端介面開發：'
    '學習Flask框架建立RESTful API，以及HTML/CSS/JavaScript前端介面設計，'
    '實現即時MJPEG影像串流、狀態更新、虛擬搖桿操控等功能。',
    indent=True
)
body(
    '(4) 學習多模態感測器整合與融合演算法：'
    '學習音訊訊號處理（FFT頻譜分析、VAD語音偵測）、'
    'rPPG光體積描記術原理與實作、'
    '以及加權融合演算法的設計與參數調校。',
    indent=True
)

# ================================================================
# 五、預期結果
# ================================================================
heading('五、預期結果')
body(
    '本研究預期開發完成之搜救機器人系統，能夠在災後建物環境中自主巡邏探索，'
    '即時偵測受困者並評估其生命跡象狀態。具體預期成果如下：',
    indent=True
)
body(
    '（1）即時人體偵測與姿態辨識：透過Hailo-8 NPU加速YOLOv8推論，'
    '達到每秒15幀以上的即時偵測速度，能夠辨識倒地、蜷縮、疑似不適等異常姿態，'
    '偵測準確率預期達85%以上。',
    indent=True
)
body(
    '（2）非接觸式生命跡象偵測：透過rPPG技術，在受困者靜止且距離≤50cm的條件下，'
    '估算其心率數值，為搜救人員提供受困者生命跡象的初步判斷依據。',
    indent=True
)
body(
    '（3）多模態融合風險評估：整合視覺、音訊、距離、互動回應及生命跡象等六種感測模態，'
    '計算VictimScore綜合風險分數，自動區分風險等級並觸發對應搜救行動。',
    indent=True
)
body(
    '（4）自主巡邏與避障：搭載超聲波旋轉掃描策略與熱區記憶地圖，'
    '使機器人能夠有效探索未知環境、避免重複搜索，並在狹小空間中靈活避障。',
    indent=True
)
body(
    '（5）遠端監控與即時通報：透過Flask網頁介面提供即時影像串流與遙控操作，'
    '發現受困者時自動透過Telegram Bot發送照片與座標資訊至搜救指揮中心，'
    '大幅縮短資訊傳遞時間，提升搜救效率。',
    indent=True
)

# ================================================================
# 六、參考文獻
# ================================================================
heading('六、參考文獻')
refs = [
    '[1] Ultralytics (2024). YOLOv8 官方文件. Retrieved from https://docs.ultralytics.com/',
    '[2] Ultralytics (2024). YOLOv8 GitHub 原始碼. Retrieved from https://github.com/ultralytics/ultralytics',
    '[3] Hailo Technologies (2024). Hailo-8 AI Processor. Retrieved from https://hailo.ai/products/ai-accelerators/hailo-8-ai-accelerator/',
    '[4] Raspberry Pi Foundation (2024). Raspberry Pi 5 官方規格. Retrieved from https://www.raspberrypi.com/products/raspberry-pi-5/',
    '[5] iPhysioMeter (2023). Remote Photoplethysmography (rPPG) 技術介紹. Retrieved from https://iphysiometer.com/rppg/',
    '[6] RobotShop (2024). Mecanum Wheels - How They Work. Retrieved from https://www.robotshop.com/community/tutorials/show/mecanum-wheels-how-they-work',
    '[7] Flask 官方文件 (2024). Flask Web Framework. Retrieved from https://flask.palletsprojects.com/',
    '[8] gpiozero 官方文件 (2024). GPIO Zero Documentation. Retrieved from https://gpiozero.readthedocs.io/',
    '[9] 內政部消防署 (2024). 消防統計. Retrieved from https://www.nfa.gov.tw/',
]
for ref in refs:
    body(ref)

# ================================================================
# 七、需要指導教授指導內容
# ================================================================
heading('七、需要指導教授指導內容')
body('7.1\tHailo-8 NPU模型編譯與優化部署')
body('7.2\trPPG生命跡象偵測演算法之參數校準與驗證')
body('7.3\t多模態融合權重之最佳化調校')
body('7.4\t自主巡邏策略之可靠性驗證與改進')

# ── 儲存 ──
out = '/Users/user/Desktop/具生命跡象偵測之搜救機器人系統/output/專題內容_一至五.docx'
doc.save(out)
print(f'已儲存: {out}')
